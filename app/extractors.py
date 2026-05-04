from __future__ import annotations

from dataclasses import dataclass
import logging
from pathlib import Path
import re
from typing import Iterable

import docx
from pypdf import PdfReader

LOGGER = logging.getLogger(__name__)


@dataclass(frozen=True)
class TextChunk:
    location: str
    text: str


@dataclass(frozen=True)
class PositionedLine:
    location: str
    text: str
    indent: float


def _group_lines(words: list[dict], *, y_tolerance: float = 3.0) -> list[str]:
    lines: list[list[dict]] = []
    for word in sorted(words, key=lambda w: (float(w["top"]), float(w["x0"]))):
        if not lines:
            lines.append([word])
            continue
        prev_top = float(lines[-1][0]["top"])
        if abs(float(word["top"]) - prev_top) <= y_tolerance:
            lines[-1].append(word)
        else:
            lines.append([word])

    rendered: list[str] = []
    for line_words in lines:
        parts = [str(w["text"]).strip() for w in sorted(line_words, key=lambda w: float(w["x0"]))]
        line = " ".join(part for part in parts if part)
        if line:
            rendered.append(line)
    return rendered


def _group_positioned_lines(
    words: list[dict],
    *,
    location: str,
    base_x: float = 0.0,
    y_tolerance: float = 3.0,
) -> list[PositionedLine]:
    lines: list[list[dict]] = []
    for word in sorted(words, key=lambda w: (float(w["top"]), float(w["x0"]))):
        if not lines:
            lines.append([word])
            continue
        prev_top = float(lines[-1][0]["top"])
        if abs(float(word["top"]) - prev_top) <= y_tolerance:
            lines[-1].append(word)
        else:
            lines.append([word])

    rendered: list[PositionedLine] = []
    for line_words in lines:
        ordered = sorted(line_words, key=lambda w: float(w["x0"]))
        text = " ".join(str(w["text"]).strip() for w in ordered if str(w["text"]).strip())
        if text:
            rendered.append(
                PositionedLine(
                    location=location,
                    text=text,
                    indent=float(ordered[0]["x0"]) - base_x,
                )
            )
    return rendered


def _split_columns(words: list[dict], *, depth: int = 2) -> list[list[dict]]:
    if depth <= 0 or len(words) < 40:
        return [words]

    xs = sorted(float(w["x0"]) for w in words)
    if len(xs) < 2:
        return [words]

    max_gap = 0.0
    split_idx = -1
    for i in range(len(xs) - 1):
        gap = xs[i + 1] - xs[i]
        if gap > max_gap:
            max_gap = gap
            split_idx = i

    x_span = max(xs) - min(xs)
    gap_threshold = max(40.0, x_span * 0.2)
    if split_idx < 0 or max_gap < gap_threshold:
        return [words]

    split_x = (xs[split_idx] + xs[split_idx + 1]) / 2.0
    left = [w for w in words if float(w["x0"]) <= split_x]
    right = [w for w in words if float(w["x0"]) > split_x]

    if len(left) < 18 or len(right) < 18:
        return [words]

    return _split_columns(left, depth=depth - 1) + _split_columns(right, depth=depth - 1)


def _clean_joined_text(parts: list[str]) -> str:
    text = " ".join(part.strip() for part in parts if part.strip())
    replacements = {
        "\x00": "ti",
        "\ufb00": "ff",
        "\ufb01": "fi",
        "\ufb02": "fl",
        "\ufb03": "ffi",
        "\ufb04": "ffl",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Drop typical PDF glyph-noise sequences, but keep isolated tokens that
    # could be meaningful in document text.
    text = re.sub(r"(?:/g\d+[A-Za-z]*){2,}", " ", text)
    text = re.sub(r"(?:\s+/g\d+[A-Za-z]*){3,}", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _has_numeric_prefix(path: Path, prefix: str) -> bool:
    return re.match(rf"^{re.escape(prefix)}(?:\D|$)", path.name) is not None


def _is_us_entry_start(text: str) -> bool:
    if not text:
        return False

    continuation_prefixes = (
        "Secondary ",
        "Sanctions ",
        "Registration ",
        "Identification ",
        "Executive ",
        "Vessel ",
        "MMSI ",
        "(Linked ",
        "[",
        "Additional ",
        "Subject ",
        "Former ",
        "Other ",
        "Transactions ",
        "Prohibited ",
        "Section ",
        "alt. ",
        "March ",
        "To: ",
        "OFFICE ",
    )
    if text.startswith(continuation_prefixes):
        return False

    vessel_markers = (
        "General Cargo",
        "Crude Oil Tanker",
        "Products Tanker",
        "Bulk Carrier",
        "Container Ship",
        "Chemical/Products Tanker",
        "Hopper Barge",
        "Yacht",
        "Tanker",
        "Cargo Ship",
        "Passenger",
        "Ro-Ro",
    )
    return any(marker in text for marker in vessel_markers)


def _extract_us_sdn_pdf(path: Path) -> Iterable[TextChunk]:
    import pdfplumber

    header_re = re.compile(
        r"^(OFFICE OF FOREIGN ASSETS CONTROL|March \d{1,2}, \d{4}|- \d+ -|To: )"
    )
    column_boxes = (
        (40, 40, 205, 740),
        (205, 40, 375, 740),
        (375, 40, 545, 740),
    )

    with pdfplumber.open(str(path)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            page_lines: list[PositionedLine] = []
            for col_idx, box in enumerate(column_boxes, start=1):
                crop = page.crop(box)
                words = crop.extract_words(
                    x_tolerance=2,
                    y_tolerance=3,
                    keep_blank_chars=False,
                    use_text_flow=False,
                ) or []
                if not words:
                    continue
                page_lines.extend(
                    _group_positioned_lines(
                        words,
                        location=f"page {page_index}",
                        base_x=box[0],
                    )
                )

            lines = [line.text.strip() for line in page_lines if line.text.strip() and not header_re.match(line.text.strip())]
            if not lines:
                continue

            start_indices = [idx for idx, text in enumerate(lines) if _is_us_entry_start(text)]
            if not start_indices:
                continue

            for idx, text in enumerate(lines):
                if "IMO " not in text:
                    continue

                start = 0
                for candidate in start_indices:
                    if candidate <= idx:
                        start = candidate
                    else:
                        break

                end = len(lines)
                for candidate in start_indices:
                    if candidate > idx:
                        end = candidate
                        break

                chunk_text = _clean_joined_text(lines[start:end])
                if chunk_text:
                    yield TextChunk(location=f"page {page_index}", text=chunk_text)


def _extract_eu_docx(path: Path) -> Iterable[TextChunk]:
    document = docx.Document(str(path))
    if not document.tables:
        return

    for row_index, row in enumerate(document.tables[0].rows, start=1):
        cells = row.cells
        if len(cells) < 3:
            continue
        entry_no = _clean_joined_text([p.text for p in cells[0].paragraphs])
        vessel_name = _clean_joined_text([p.text for p in cells[1].paragraphs])
        imo_raw = _clean_joined_text([p.text for p in cells[2].paragraphs])
        text = _clean_joined_text([p.text for p in cells[3].paragraphs]) if len(cells) > 3 else ""
        listed_at = _clean_joined_text([p.text for p in cells[4].paragraphs]) if len(cells) > 4 else ""
        imo_match = re.search(r"\b(\d{7})\b", imo_raw)
        if not imo_match:
            continue
        # EU list may store IMO with labels like "IMO number: 1234567".
        # Keep only the 7-digit value from the dedicated IMO column.
        imo = imo_match.group(1)
        row_text = " | ".join(
            part
            for part in (entry_no, vessel_name, imo, text, listed_at, f"[EU_IMO:{imo}]")
            if part
        )
        yield TextChunk(location=f"row {row_index}", text=row_text)


def _extract_uk_ship_pdf(path: Path) -> Iterable[TextChunk]:
    entry_header_re = re.compile(r"^\d+\.$")
    footer_re = re.compile(r"^Page \d+ of \d+$")
    imo_re = re.compile(r"\bIMO\s*\d{7}\b|\b\d{7}\b")
    reader = PdfReader(str(path))
    current_entry_number: str | None = None
    current_entry_page: int | None = None
    current_lines: list[str] = []

    def flush() -> TextChunk | None:
        nonlocal current_entry_number, current_entry_page, current_lines
        if not current_entry_number or current_entry_page is None or not current_lines:
            current_entry_number = None
            current_entry_page = None
            current_lines = []
            return None

        chunk = TextChunk(
            location=f"entry {current_entry_number}, page {current_entry_page}",
            text=_clean_joined_text(current_lines),
        )
        current_entry_number = None
        current_entry_page = None
        current_lines = []
        return chunk

    total_pages = len(reader.pages)
    for page_index, page in enumerate(reader.pages, start=1):
        if page_index == 1 or page_index % 50 == 0:
            LOGGER.info("UK PDF progress: %s/%s pages", page_index, total_pages)

        text = page.extract_text() or ""
        if not text.strip():
            continue

        lines = [line.replace("\x00", "ti").strip() for line in text.splitlines()]
        for line in lines:
            if not line or footer_re.match(line):
                continue

            if entry_header_re.match(line):
                chunk = flush()
                if chunk:
                    yield chunk
                current_entry_number = line[:-1]
                current_entry_page = page_index
                continue

            if current_entry_number is not None:
                current_lines.append(line)
                continue

            # Some UK list revisions may include IMO lines outside numbered
            # entries (for example in trailing sections). Keep those lines.
            if imo_re.search(line):
                yield TextChunk(location=f"page {page_index}", text=line)

    chunk = flush()
    if chunk:
        yield chunk


def extract_pdf(path: Path) -> Iterable[TextChunk]:
    if _has_numeric_prefix(path, "01"):
        yield from _extract_us_sdn_pdf(path)
        return

    if _has_numeric_prefix(path, "03"):
        yield from _extract_uk_ship_pdf(path)
        return

    try:
        import pdfplumber

        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                words = page.extract_words(
                    x_tolerance=2,
                    y_tolerance=3,
                    keep_blank_chars=False,
                    use_text_flow=False,
                ) or []

                if not words:
                    continue

                columns = _split_columns(words)
                columns = sorted(columns, key=lambda col: min(float(w["x0"]) for w in col))
                for col_idx, column_words in enumerate(columns, start=1):
                    lines = _group_lines(column_words)
                    text = "\n".join(lines).strip()
                    if not text:
                        continue
                    location = f"page {i}" if len(columns) == 1 else f"page {i} col {col_idx}"
                    yield TextChunk(location=location, text=text)
        return
    except Exception:
        pass

    reader = PdfReader(str(path))
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            yield TextChunk(location=f"page {i}", text=text)


def extract_docx(path: Path) -> Iterable[TextChunk]:
    if _has_numeric_prefix(path, "02"):
        yield from _extract_eu_docx(path)
        return

    document = docx.Document(str(path))
    for i, paragraph in enumerate(document.paragraphs, start=1):
        text = (paragraph.text or "").strip()
        if text:
            yield TextChunk(location=f"paragraph {i}", text=text)


def extract_text(path: Path) -> Iterable[TextChunk]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        yield from extract_pdf(path)
        return

    if suffix in {".docx", ".doxc"}:
        yield from extract_docx(path)
        return

    raise ValueError(f"Unsupported file type: {path.suffix}")
